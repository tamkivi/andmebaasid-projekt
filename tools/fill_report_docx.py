#!/usr/bin/env python3
from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
import re
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

from sql_ddl import SQL_DDL

ROOT = Path(__file__).resolve().parents[1]
DST = ROOT / "Jousaali_infosusteemi_treeningute_funktsionaalne_allsusteem.docx"
DIAGRAM_DIR = ROOT / "work" / "generated_diagrams"


AUTHORS = "Tristan Aik Sild, Gustav Tamkivi"
STUDY_GROUP = "IAIB23"
MATRICULATION_NUMBERS = "Tristan Aik Sild: 253782IAIB; Gustav Tamkivi: 253787IAIB"
AUTHOR_EMAILS = "Gustav Tamkivi: gustav@taltech.ee; Tristan Aik Sild: trists@taltech.ee"
SYSTEM_NAME = "Jõusaali infosüsteem"
SUBSYSTEM_NAME = "Treeningute funktsionaalne allsüsteem"
REGISTER_NAME = "Treeningute register"
UNIVERSITY = "TALLINNA TEHNIKAÜLIKOOL"
FACULTY = "Infotehnoloogia teaduskond"
INSTITUTE = "Tarkvarateaduse instituut"
COURSE = "Andmebaasid I, ITI0206"
SUPERVISOR = "Erki Eessaar"


ACTORS = [
    ("Treener", "Sisemine pädevusala", "Registreerib, muudab ja aktiveerib treeninguid."),
    ("Juhataja", "Sisemine pädevusala", "Jälgib treeningute portfelli ning lõpetab treeninguid."),
    ("Klient", "Väline pädevusala", "Soovib leida sobiva aktiivse treeningu."),
    ("Uudistaja", "Väline pädevusala", "Soovib vaadata avalikult nähtavat treeningute infot."),
    ("Klassifikaatorite haldur", "Sisemine pädevusala", "Haldab klassifikaatoreid, sh treeningu seisundeid ja kategooriaid."),
    ("Töötajate haldur", "Sisemine pädevusala", "Haldab töötajate ja rollide põhiandmeid."),
]


MAIN_OBJECTS = [
    ("Isik", "Jõusaali kliendi, töötaja või süsteemi kasutaja üldandmed."),
    ("Töötaja", "Jõusaali heaks töötav isik ja tema rollid."),
    ("Treening", "Jõusaalis pakutav treeninguteenus."),
    ("Klassifikaator", "Ühiselt hallatavad seisundid, tüübid ja kategooriad."),
]


SUBSYSTEMS_REGISTERS = [
    ("Isikute funktsionaalne allsüsteem", "Isikute register", "Isik"),
    ("Töötajate funktsionaalne allsüsteem", "Töötajate register", "Töötaja"),
    ("Treeningute funktsionaalne allsüsteem", "Treeningute register", "Treening"),
    ("Klassifikaatorite funktsionaalne allsüsteem", "Klassifikaatorite register", "Klassifikaator"),
]


STATEMENTS = [
    "Isik esitab jõusaali kliendi, töötaja või süsteemi kasutaja üldandmeid.",
    "Treening on jõusaalis pakutav teenus.",
    "Treener registreerib treeningu.",
    "Juhataja lõpetab treeningu.",
    "Klient vaatab aktiivse treeningu avalikke andmeid.",
    "Uudistaja vaatab avalikult nähtavaid aktiivseid treeninguid.",
    "Klassifikaatorite haldur haldab klassifikaatoreid.",
    "Töötajate haldur haldab töötajaid.",
    "Treeningul on üks seisund.",
    "Treening kuulub nulli või mitmesse treeningu kategooriasse.",
]


PROCESSES = [
    ("Kasutaja tuvastamine", "Isik soovib kasutada tööalast süsteemifunktsiooni.", "Isiku kasutajakonto ja töötaja pädevusala on tuvastatud."),
    ("Treeningu registreerimine", "Treenerist töötaja saab teabe uue treeningu kohta.", "Treening on ootel seisundis registreeritud ja klassifikaatoritega seotud."),
    ("Treeningu muutmine", "Treeningu andmetes ilmneb viga või muutus.", "Treeningu andmed on ajakohased."),
    ("Treeningu aktiveerimine", "Treening on valmis klientidele nähtavaks tegemiseks.", "Treening on aktiivne."),
    ("Treeningu ajutiselt kasutusest eemaldamine", "Treeninguga ilmneb ajutine probleem.", "Treening on mitteaktiivne."),
    ("Treeningu lõpetamine", "Treeningut ei pakuta enam või probleem on püsiv.", "Treening on lõppenud."),
    ("Treeningute vaatamine", "Klient või uudistaja soovib treeningute infot.", "Kasutaja näeb lubatud treeninguandmeid."),
]


TRAINING_LIST_DETAILS = "iga rea kohta treeningu kood, nimetus, kategooria, seisund, kestus minutites, hind, maksimaalne osalejate arv, registreerimise aeg ja viimase muutmise aeg"
PUBLIC_TRAINING_LIST_DETAILS = "iga rea kohta treeningu kood, nimetus, kategooria, seisund, kestus minutites, hind, maksimaalne osalejate arv ja vajalik varustus"
CATEGORY_LIST_DETAILS = "iga rea kohta kategooria kood, nimetus, tüüp ja aktiivsuse tunnus"


HIGH_LEVEL_USE_CASES = [
    ("Tuvasta kasutaja", "Treener, Juhataja, Klassifikaatorite haldur, Töötajate haldur", "Kasutaja esitab süsteemile enda tuvastamiseks vajalikud andmed ning süsteem seob kasutaja tema pädevusalaga."),
    ("Registreeri treening", "Treener", "Treener registreerib uue treeningu põhiandmed ja seob treeningu sobivate kategooriatega."),
    ("Muuda treeningut", "Treener", "Treener muudab ootel või mitteaktiivse treeningu põhiandmeid ja kategooriaseoseid."),
    ("Unusta treening", "Treener", "Treener eemaldab ootel treeningu edasisest kasutusest, kui treeningut ei ole vaja pakkuda."),
    ("Aktiveeri treening", "Treener", "Treener muudab ootel või mitteaktiivse treeningu aktiivseks, kui treening kuulub vähemalt ühte kategooriasse."),
    ("Muuda treening mitteaktiivseks", "Treener", "Treener eemaldab aktiivse treeningu ajutiselt kasutusest, kui treeninguga on ilmnenud ajutine probleem."),
    ("Lõpeta treening", "Juhataja", "Juhataja lõpetab aktiivse või mitteaktiivse treeningu, kui treeningut enam ei pakuta."),
    ("Vaata aktiivseid treeninguid", "Klient, Uudistaja", "Klient või uudistaja vaatab aktiivsete treeningute avalikke andmeid."),
    ("Vaata kõiki treeninguid", "Juhataja", "Juhataja vaatab kõigi treeningute tööalaseid andmeid ja seisundeid."),
    ("Vaata kõiki ootel või mitteaktiivseid treeninguid", "Treener", "Treener vaatab ootel ja mitteaktiivseid treeninguid, et valida muutmiseks või aktiveerimiseks sobiv treening."),
    ("Vaata treeningute koondaruannet", "Juhataja", "Juhataja vaatab treeningute arvu seisundite ja kategooriate kaupa."),
]


OPERATIONS = [
    ("OP3", "Registreeri treening(\n    p_treeningu_kood,\n    p_nimetus,\n    p_kirjeldus,\n    p_kestus_minutites,\n    p_maksimaalne_osalejate_arv,\n    p_vajalik_varustus,\n    p_hind,\n    p_treeningu_kategooria_koodid,\n    p_registreerija_e_meil\n)", [
        "Treener on autenditud ja autoriseeritud.",
        "Treeningu_seisundi_liik eksemplar tsl (mille kood on Ootel) on registreeritud.",
        "Kõik p_treeningu_kategooria_koodid hulka kuuluvad Treeningu_kategooria eksemplarid on registreeritud.",
        "Kasutajakonto eksemplar k (mille e_meil = p_registreerija_e_meil) on registreeritud."
    ], [
        "Treening eksemplar t on registreeritud.",
        "t.treeningu_kood := p_treeningu_kood.",
        "t.nimetus := p_nimetus.",
        "t.kirjeldus := p_kirjeldus.",
        "t.kestus_minutites := p_kestus_minutites.",
        "t.maksimaalne_osalejate_arv := p_maksimaalne_osalejate_arv.",
        "t.vajalik_varustus := p_vajalik_varustus.",
        "t.hind := p_hind.",
        "t.seisund := Ootel.",
        "t.registreerija_e_meil := p_registreerija_e_meil.",
        "t.viimase_muutja_e_meil := p_registreerija_e_meil.",
        "t.reg_aeg := hetke kuupäev ja kellaaeg.",
        "t.viimase_muutm_aeg := hetke kuupäev ja kellaaeg.",
        "t ja kõik p_treeningu_kategooria_koodid järgi leitud Treeningu_kategooria eksemplarid on seotud."
    ]),
    ("OP6", "Unusta treening(\n    p_treeningu_kood,\n    p_muutja_e_meil\n)", [
        "Treening eksemplar t (mille treeningu_kood = p_treeningu_kood) on ootel seisundis registreeritud.",
        "Kasutajakonto eksemplar k (mille e_meil = p_muutja_e_meil) on registreeritud."
    ], [
        "t.seisund := Unustatud.",
        "t.viimase_muutja_e_meil := p_muutja_e_meil.",
        "t.viimase_muutm_aeg := hetke kuupäev ja kellaaeg."
    ]),
    ("OP8", "Muuda treeningut(\n    p_treeningu_kood,\n    p_nimetus,\n    p_kirjeldus,\n    p_kestus_minutites,\n    p_maksimaalne_osalejate_arv,\n    p_vajalik_varustus,\n    p_hind,\n    p_muutja_e_meil\n)", [
        "Treening eksemplar t (mille treeningu_kood = p_treeningu_kood) on ootel või mitteaktiivses seisundis registreeritud.",
        "Kasutajakonto eksemplar k (mille e_meil = p_muutja_e_meil) on registreeritud."
    ], [
        "t.nimetus := p_nimetus.",
        "t.kirjeldus := p_kirjeldus.",
        "t.kestus_minutites := p_kestus_minutites.",
        "t.maksimaalne_osalejate_arv := p_maksimaalne_osalejate_arv.",
        "t.vajalik_varustus := p_vajalik_varustus.",
        "t.hind := p_hind.",
        "t.viimase_muutja_e_meil := p_muutja_e_meil.",
        "t.viimase_muutm_aeg := hetke kuupäev ja kellaaeg."
    ]),
    ("OP9", "Lisa treeningu kategooria seos(\n    p_treeningu_kood,\n    p_treeningu_kategooria_kood\n)", [
        "Treening eksemplar t (mille treeningu_kood = p_treeningu_kood) on registreeritud.",
        "Treeningu_kategooria eksemplar tk (mille kood = p_treeningu_kategooria_kood) on registreeritud.",
        "t ja tk seos ei ole registreeritud."
    ], [
        "t ja tk seos on registreeritud."
    ]),
    ("OP10", "Eemalda treeningu kategooria seos(\n    p_treeningu_kood,\n    p_treeningu_kategooria_kood\n)", [
        "Treening eksemplar t (mille treeningu_kood = p_treeningu_kood) on registreeritud.",
        "Treeningu_kategooria eksemplar tk (mille kood = p_treeningu_kategooria_kood) on registreeritud.",
        "t ja tk seos on registreeritud."
    ], [
        "t ja tk seos on kustutatud."
    ]),
    ("OP11", "Aktiveeri treening(\n    p_treeningu_kood,\n    p_muutja_e_meil\n)", [
        "Treening eksemplar t (mille treeningu_kood = p_treeningu_kood) on ootel või mitteaktiivses seisundis registreeritud.",
        "t on seotud vähemalt ühe Treeningu_kategooria eksemplariga.",
        "Kasutajakonto eksemplar k (mille e_meil = p_muutja_e_meil) on registreeritud."
    ], [
        "t.seisund := Aktiivne.",
        "t.viimase_muutja_e_meil := p_muutja_e_meil.",
        "t.viimase_muutm_aeg := hetke kuupäev ja kellaaeg."
    ]),
    ("OP13", "Muuda treening mitteaktiivseks(\n    p_treeningu_kood,\n    p_muutja_e_meil\n)", [
        "Treening eksemplar t (mille treeningu_kood = p_treeningu_kood) on aktiivses seisundis registreeritud.",
        "Kasutajakonto eksemplar k (mille e_meil = p_muutja_e_meil) on registreeritud."
    ], [
        "t.seisund := Mitteaktiivne.",
        "t.viimase_muutja_e_meil := p_muutja_e_meil.",
        "t.viimase_muutm_aeg := hetke kuupäev ja kellaaeg."
    ]),
    ("OP15", "Lõpeta treening(\n    p_treeningu_kood,\n    p_muutja_e_meil\n)", [
        "Treening eksemplar t (mille treeningu_kood = p_treeningu_kood) on aktiivses või mitteaktiivses seisundis registreeritud.",
        "Kasutajakonto eksemplar k (mille e_meil = p_muutja_e_meil) on registreeritud."
    ], [
        "t.seisund := Lõppenud.",
        "t.viimase_muutja_e_meil := p_muutja_e_meil.",
        "t.viimase_muutm_aeg := hetke kuupäev ja kellaaeg."
    ]),
]


EXTENDED_USE_CASES = [
    {
        "name": "Tuvasta kasutaja",
        "actor": "Treener, Juhataja, Klassifikaatorite haldur, Töötajate haldur",
        "stakeholders": [
            ("Treener", "Soovib pääseda enda tööülesannete täitmiseks lubatud kasutusjuhtudeni."),
            ("Juhataja", "Soovib, et süsteemi tööalaseid funktsioone kasutaksid ainult õigustatud kasutajad."),
            ("Klassifikaatorite haldur", "Soovib pääseda klassifikaatorite haldamiseks lubatud tööalaste funktsioonideni."),
            ("Töötajate haldur", "Soovib pääseda töötajate andmete haldamiseks lubatud tööalaste funktsioonideni."),
        ],
        "trigger": "Kasutaja soovib alustada tööalase süsteemifunktsiooni kasutamist.",
        "pre": ["Kasutajakonto on registreeritud ja aktiivne."],
        "post": ["Kasutaja on autenditud ja tema pädevusalad on tuvastatud."],
        "steps": [
            ("Treener, juhataja, klassifikaatorite haldur või töötajate haldur avaldab soovi süsteemi siseneda.", ""),
            ("Süsteem kuvab tuvastamiseks vajaliku vormi.", ""),
            ("Treener, juhataja, klassifikaatorite haldur või töötajate haldur sisestab e-posti aadressi ja parooli.", ""),
            ("Süsteem tuvastab kasutaja ning leiab tema pädevusalad.", ""),
        ],
        "extensions": [
            "4a. Kui kasutajakontot ei leita või parool ei vasta kontole, siis kasutajat ei autendita.",
            "4b. Kui kasutajakonto ei ole aktiivne, siis kasutajat ei autendita.",
        ],
    },
    {
        "name": "Registreeri treening",
        "actor": "Treener",
        "stakeholders": [
            ("Treener", "Soovib uue treeningu andmed kiiresti ja korrektselt registreerida."),
            ("Juhataja", "Soovib, et treeningute portfell oleks ajakohane."),
            ("Klient", "Soovib hiljem näha terviklikku infot aktiivsete treeningute kohta."),
        ],
        "trigger": "Treener saab teabe uue jõusaalis pakutava treeningu kohta.",
        "pre": ["Treener on autenditud ja autoriseeritud.", "Treeningu seisundi ja kategooria klassifikaatorid on registreeritud."],
        "post": ["Treening on registreeritud ootel seisundis.", "Treening on seotud valitud treeningu kategooriatega."],
        "steps": [
            ("Treener avaldab soovi registreerida uus treening.", ""),
            (f"Süsteem kuvab treeningu sisestamiseks vajaliku vormi ning aktiivsed treeningu kategooriad; kategooriate nimekirjas kuvatakse {CATEGORY_LIST_DETAILS}.", ""),
            ("Treener sisestab treeningu nimetuse, kirjelduse, kestuse, maksimaalse osalejate arvu, vajaliku varustuse, hinna ja kategooriad.", ""),
            ("Süsteem kontrollib sisestatud andmeid.", ""),
            ("Süsteem salvestab treeningu andmed ja kategooriaseosed.", "OP3"),
            (f"Süsteem kuvab registreeritud treeningu detailandmed: {TRAINING_LIST_DETAILS}.", ""),
        ],
        "extensions": [
            "4a. Kui kohustuslik väärtus puudub või ei vasta piirangule, siis salvestamine ebaõnnestub.",
            "5a. Kui sama nimetusega treening on juba registreeritud, siis salvestamine ebaõnnestub.",
        ],
    },
    {
        "name": "Muuda treeningut",
        "actor": "Treener",
        "stakeholders": [
            ("Treener", "Soovib parandada treeningu andmeid või ajakohastada kategooriaseoseid."),
            ("Juhataja", "Soovib, et treeningute register sisaldaks korrektseid tööandmeid."),
            ("Klient", "Soovib, et aktiivseks muutuvad treeningud oleksid sisuliselt õiged."),
        ],
        "trigger": "Treener märkab ootel või mitteaktiivse treeningu andmetes muutmise vajadust.",
        "pre": ["Treener on autenditud ja autoriseeritud.", "Treening on ootel või mitteaktiivses seisundis."],
        "post": ["Treeningu põhiandmed ja kategooriaseosed on ajakohased."],
        "steps": [
            ("Treener avaldab soovi muuta treeningu andmeid.", ""),
            (f"Süsteem kuvab ootel ja mitteaktiivsed treeningud; nimekirjas kuvatakse {TRAINING_LIST_DETAILS}.", ""),
            ("Treener valib muudetava treeningu.", ""),
            (f"Süsteem kuvab valitud treeningu detailandmed ja kategooriaseosed: {TRAINING_LIST_DETAILS}.", ""),
            ("Treener muudab treeningu põhiandmeid ning lisab või eemaldab kategooriaseoseid.", ""),
            ("Süsteem salvestab treeningu põhiandmete muudatused.", "OP8"),
            ("Süsteem salvestab lisatud kategooriaseosed.", "OP9"),
            ("Süsteem kustutab eemaldatud kategooriaseosed.", "OP10"),
        ],
        "extensions": [
            "2a. Kui süsteemis ei ole ühtegi ootel või mitteaktiivset treeningut, siis ei saa Treener jätkata.",
            "6a. Kui muudetud andmed ei vasta piirangutele, siis muudatuste salvestamine ebaõnnestub.",
        ],
    },
    {
        "name": "Unusta treening",
        "actor": "Treener",
        "stakeholders": [
            ("Treener", "Soovib eemaldada ootel treeningu edasisest töövoost."),
            ("Juhataja", "Soovib, et treeningute portfell ei sisaldaks realiseerumata ettepanekuid."),
        ],
        "trigger": "Treener saab teabe, et ootel treeningut ei hakata pakkuma.",
        "pre": ["Treener on autenditud ja autoriseeritud.", "Treening on ootel seisundis."],
        "post": ["Treening on unustatud seisundis."],
        "steps": [
            ("Treener avaldab soovi unustada treening.", ""),
            (f"Süsteem kuvab ootel treeningud; nimekirjas kuvatakse {TRAINING_LIST_DETAILS}.", ""),
            ("Treener valib unustatava treeningu.", ""),
            ("Süsteem salvestab treeningu seisundimuudatuse.", "OP6"),
            ("Süsteem eemaldab unustatud treeningu ootel treeningute nimekirjast.", ""),
        ],
        "extensions": [
            "2a. Kui süsteemis ei ole ühtegi ootel treeningut, siis ei saa Treener treeningut unustada.",
        ],
    },
    {
        "name": "Aktiveeri treening",
        "actor": "Treener",
        "stakeholders": [
            ("Treener", "Soovib valmis treeningu klientidele nähtavaks teha."),
            ("Klient", "Soovib näha ainult pakutavaid ja osalemiseks sobivaid treeninguid."),
            ("Juhataja", "Soovib, et avalikus vaates oleks ainult kontrollitud treeningud."),
        ],
        "trigger": "Treener otsustab, et ootel või mitteaktiivne treening on valmis aktiivseks muutmiseks.",
        "pre": ["Treener on autenditud ja autoriseeritud.", "Treening on ootel või mitteaktiivses seisundis."],
        "post": ["Treening on aktiivses seisundis."],
        "steps": [
            ("Treener avaldab soovi aktiveerida treening.", ""),
            (f"Süsteem kuvab ootel ja mitteaktiivsed treeningud; nimekirjas kuvatakse {TRAINING_LIST_DETAILS}.", ""),
            ("Treener valib nimekirjast treeningu.", ""),
            ("Süsteem kontrollib, et treening kuulub vähemalt ühte treeningu kategooriasse.", ""),
            ("Süsteem salvestab treeningu seisundimuudatuse.", "OP11"),
            (f"Süsteem kuvab muudetud treeningu detailandmed: {TRAINING_LIST_DETAILS}.", ""),
        ],
        "extensions": [
            "3a. Kui nimekirjas ei ole ühtegi ootel või mitteaktiivset treeningut, siis ei saa Treener jätkata.",
            "4a. Kui treening ei kuulu ühtegi treeningu kategooriasse, siis aktiveerimine ebaõnnestub.",
        ],
    },
    {
        "name": "Muuda treening mitteaktiivseks",
        "actor": "Treener",
        "stakeholders": [
            ("Treener", "Soovib ajutise probleemiga treeningu klientidele nähtavast pakkumisest eemaldada."),
            ("Klient", "Soovib näha ainult tegelikult pakutavaid aktiivseid treeninguid."),
            ("Juhataja", "Soovib, et treeningute avalik pakkumine oleks kontrollitud."),
        ],
        "trigger": "Treener saab teabe aktiivse treeningu ajutisest probleemist.",
        "pre": ["Treener on autenditud ja autoriseeritud.", "Treening on aktiivses seisundis."],
        "post": ["Treening on mitteaktiivses seisundis."],
        "steps": [
            ("Treener avaldab soovi muuta treening mitteaktiivseks.", ""),
            (f"Süsteem kuvab aktiivsed treeningud; nimekirjas kuvatakse {TRAINING_LIST_DETAILS}.", ""),
            ("Treener valib mitteaktiivseks muudetava treeningu.", ""),
            ("Süsteem salvestab treeningu seisundimuudatuse.", "OP13"),
            ("Süsteem eemaldab treeningu aktiivsete treeningute avalikust nimekirjast.", ""),
        ],
        "extensions": [
            "2a. Kui süsteemis ei ole ühtegi aktiivset treeningut, siis ei saa Treener treeningut mitteaktiivseks muuta.",
        ],
    },
    {
        "name": "Lõpeta treening",
        "actor": "Juhataja",
        "stakeholders": [
            ("Juhataja", "Soovib eemaldada püsivalt treeningu, mida jõusaal enam ei paku."),
            ("Treener", "Soovib, et treeningute nimekiri ei sisaldaks kasutusest väljas teenuseid."),
            ("Klient", "Soovib näha ainult tegelikult pakutavaid aktiivseid treeninguid."),
        ],
        "trigger": "Juhataja saab teabe, et treeningut enam ei pakuta või selle probleem on püsiv.",
        "pre": ["Juhataja on autenditud ja autoriseeritud.", "Treening on aktiivses või mitteaktiivses seisundis."],
        "post": ["Treening on lõppenud seisundis."],
        "steps": [
            ("Juhataja avaldab soovi lõpetada treening.", ""),
            (f"Süsteem kuvab kõik treeningud koos seisunditega; nimekirjas kuvatakse {TRAINING_LIST_DETAILS}.", ""),
            ("Juhataja valib lõpetatava treeningu.", ""),
            ("Süsteem salvestab treeningu seisundimuudatuse.", "OP15"),
            (f"Süsteem kuvab lõpetatud treeningu detailandmed: {TRAINING_LIST_DETAILS}.", ""),
        ],
        "extensions": [
            "2a. Kui süsteemis ei ole ühtegi aktiivset ega mitteaktiivset treeningut, siis ei saa Juhataja treeningut lõpetada.",
        ],
    },
    {
        "name": "Vaata aktiivseid treeninguid",
        "actor": "Klient, Uudistaja",
        "stakeholders": [
            ("Klient", "Soovib leida sobiva aktiivse treeningu."),
            ("Uudistaja", "Soovib näha jõusaali avalikku treeninguvalikut."),
            ("Juhataja", "Soovib, et avalikult nähtav treeninguinfo oleks ajakohane."),
        ],
        "trigger": "Klient või uudistaja soovib vaadata pakutavaid treeninguid.",
        "pre": ["Aktiivsete treeningute avalik vaade on kasutatav."],
        "post": ["Aktiivsete treeningute lubatud andmed on kasutajale kuvatud."],
        "steps": [
            ("Klient või uudistaja avaldab soovi vaadata aktiivseid treeninguid.", ""),
            (f"Süsteem kuvab aktiivsete treeningute nimekirja; nimekirjas kuvatakse {PUBLIC_TRAINING_LIST_DETAILS}.", ""),
            ("Klient või uudistaja valib treeningu, mille andmeid ta soovib täpsemalt vaadata.", ""),
            (f"Süsteem kuvab valitud aktiivse treeningu lubatud detailandmed: {PUBLIC_TRAINING_LIST_DETAILS}.", ""),
        ],
        "extensions": [
            "2a. Kui aktiivseid treeninguid ei ole, siis kuvab süsteem tühja nimekirja.",
        ],
    },
    {
        "name": "Vaata kõiki treeninguid",
        "actor": "Juhataja",
        "stakeholders": [
            ("Juhataja", "Soovib saada terviklikku tööalast ülevaadet kõigist treeningutest ja nende seisunditest."),
            ("Treener", "Soovib, et juhatajal oleks otsuste tegemiseks ajakohane ülevaade treeningute portfellist."),
        ],
        "trigger": "Juhataja soovib vaadata treeningute tööalast terviknimekirja.",
        "pre": ["Juhataja on autenditud ja autoriseeritud."],
        "post": ["Kõigi treeningute tööalane nimekiri on Juhatajale kuvatud."],
        "steps": [
            ("Juhataja avaldab soovi vaadata kõiki treeninguid.", ""),
            (f"Süsteem kuvab kõik treeningud; nimekirjas kuvatakse {TRAINING_LIST_DETAILS}.", ""),
            ("Juhataja valib treeningu, mille tööalaseid detailandmeid ta soovib vaadata.", ""),
            (f"Süsteem kuvab valitud treeningu tööalased detailandmed: {TRAINING_LIST_DETAILS}.", ""),
        ],
        "extensions": [
            "2a. Kui süsteemis ei ole ühtegi treeningut, siis kuvab süsteem tühja nimekirja.",
        ],
    },
    {
        "name": "Vaata kõiki ootel või mitteaktiivseid treeninguid",
        "actor": "Treener",
        "stakeholders": [
            ("Treener", "Soovib leida muutmiseks või aktiveerimiseks sobiva ootel või mitteaktiivse treeningu."),
            ("Juhataja", "Soovib, et treenerid töötaksid ainult treeningutega, mida on lubatud muuta või aktiveerida."),
        ],
        "trigger": "Treener soovib valida muutmiseks või aktiveerimiseks sobiva treeningu.",
        "pre": ["Treener on autenditud ja autoriseeritud."],
        "post": ["Ootel ja mitteaktiivsete treeningute nimekiri on Treenerile kuvatud."],
        "steps": [
            ("Treener avaldab soovi vaadata ootel ja mitteaktiivseid treeninguid.", ""),
            (f"Süsteem kuvab ootel ja mitteaktiivsed treeningud; nimekirjas kuvatakse {TRAINING_LIST_DETAILS}.", ""),
            ("Treener valib nimekirjast treeningu või lõpetab nimekirja vaatamise.", ""),
            (f"Süsteem kuvab valitud treeningu detailandmed: {TRAINING_LIST_DETAILS}.", ""),
        ],
        "extensions": [
            "2a. Kui süsteemis ei ole ühtegi ootel või mitteaktiivset treeningut, siis kuvab süsteem tühja nimekirja.",
        ],
    },
    {
        "name": "Vaata treeningute koondaruannet",
        "actor": "Juhataja",
        "stakeholders": [
            ("Juhataja", "Soovib saada ülevaadet treeningute jaotusest seisundite ning kategooriate kaupa."),
            ("Treener", "Soovib, et juhtimisotsused põhineksid ajakohastel andmetel."),
        ],
        "trigger": "Juhataja soovib hinnata treeningute portfelli hetkeseisu.",
        "pre": ["Juhataja on autenditud ja autoriseeritud."],
        "post": ["Koondaruanne on Juhatajale kuvatud."],
        "steps": [
            ("Juhataja avaldab soovi vaadata treeningute koondaruannet.", ""),
            ("Süsteem kuvab treeningute arvu seisundite ja kategooriate kaupa.", ""),
        ],
        "extensions": [
            "2a. Kui aruande aluseks olevaid treeninguid ei ole, siis kuvab süsteem nullväärtustega koondvaate.",
        ],
    },
]


ENTITIES = [
    ("Klassifikaator", "Üldine liigitav väärtus, mille abil kirjeldatakse süsteemis kasutatavaid seisundeid, tüüpe ja kategooriaid.", "kood"),
    ("Riik", "Klassifikaator, mis tähistab isiku identifikaatori väljastanud riiki, näiteks Eesti või Soome.", "kood"),
    ("Isiku_seisundi_liik", "Klassifikaator, mis kirjeldab isiku kasutatavust süsteemis, näiteks Aktiivne või Mitteaktiivne.", "kood"),
    ("Töötaja_seisundi_liik", "Klassifikaator, mis kirjeldab töötaja töösuhte või kasutusõiguse seisundit, näiteks Tööl või Lahkunud.", "kood"),
    ("Töötaja_roll", "Klassifikaator, mis kirjeldab töötajale antavat tööalast rolli, näiteks Treener või Juhataja.", "kood"),
    ("Treeningu_seisundi_liik", "Klassifikaator, mis määrab treeningu elutsükli lubatud seisundid, näiteks Ootel, Aktiivne, Mitteaktiivne, Lõppenud või Unustatud.", "kood"),
    ("Treeningu_kategooria_tüüp", "Klassifikaator, mis rühmitab treeningu kategooriaid, näiteks Treeningu eesmärk või Treeningu intensiivsus.", "kood"),
    ("Treeningu_kategooria", "Klassifikaator, millega treeninguid sisuliselt rühmitatakse, näiteks Jooga, Kardio või Jõutreening.", "kood"),
    ("Isik", "Jõusaaliga seotud füüsiline inimene, kes võib olla klient, töötaja või tööalase süsteemi kasutaja.", "isikukood + riigi_kood"),
    ("Kasutajakonto", "Isikuga seotud autentimisandmete kogum, mille abil lubatakse tööalaseid süsteemifunktsioone kasutada.", "e_meil"),
    ("Töötaja", "Jõusaali heaks töötav isik, kellele saab määrata tööalaseid rolle ja kelle tegevused võivad muuta treeningute registri andmeid.", "e_meil"),
    ("Töötaja_rolli_omamine", "Seos töötaja ja töötaja rolli vahel koos rolli kehtivusajaga.", "tootaja_e_meil + tootaja_roll_kood + alguse_aeg"),
    ("Treening", "Jõusaali poolt pakutav treeninguteenus, mille kohta hoitakse avaldamiseks, muutmiseks ja aruandluseks vajalikke andmeid.", "treeningu_kood"),
    ("Treeningu_kategooria_omamine", "Seos treeningu ja treeningu kategooria vahel, mis näitab, millistesse sisulistesse rühmadesse treening kuulub.", "treeningu_kood + treeningu_kategooria_kood"),
]


ATTRIBUTES = [
    ("Klassifikaator", "kood", "Klassifikaatori väärtust esitav kood {Kohustuslik. Ei tohi olla tühi string. Klassifikaatori tüübi piires unikaalne.}. Näiteväärtus: AKTIIVNE"),
    ("Klassifikaator", "nimetus", "Klassifikaatori väärtuse ametlik eestikeelne nimetus {Kohustuslik. Ei tohi olla tühi string.}. Näiteväärtus: Aktiivne"),
    ("Klassifikaator", "on_aktiivne", "Tunnus, mis näitab, kas klassifikaatori väärtust võib uute andmete liigitamisel kasutada {Kohustuslik. Lubatud väärtused on tõene ja väär.}. Näiteväärtus: tõene"),
    ("Isik", "isikukood", "Riigi poolt väljastatud isiku identifikaator {Kohustuslik. Koos riigi koodiga unikaalne. Lubatud on tähed, numbrid, tühikud, sidekriipsud, plussmärgid, võrdusmärgid ja kaldkriipsud.}. Näiteväärtus: 50101010011"),
    ("Isik", "eesnimi", "Isiku eesnimi {Võib puududa ainult siis, kui perenimi on registreeritud. Kui väärtus registreeritakse, siis ei tohi see olla tühi string.}. Näiteväärtus: Mari"),
    ("Isik", "perenimi", "Isiku perekonnanimi {Võib puududa ainult siis, kui eesnimi on registreeritud. Kui väärtus registreeritakse, siis ei tohi see olla tühi string.}. Näiteväärtus: Tamm"),
    ("Isik", "synni_kp", "Isiku sünnikuupäev {Kohustuslik. Lubatud vahemik on 01.01.1900 kuni 31.12.2100.}. Näiteväärtus: 12.05.1995"),
    ("Isik", "elukoht", "Isiku alalise elukoha aadress {Võib puududa. Kui väärtus registreeritakse, siis ei tohi see olla tühi string.}. Näiteväärtus: Akadeemia tee 5, Tallinn"),
    ("Isik", "e_meil", "Isiku e-posti aadress, mida kasutatakse temaga elektrooniliseks suhtlemiseks ja kasutajakonto sidumiseks {Kohustuslik. Peab sisaldama märki @. Võib olla kuni 254 märki pikk. Unikaalne.}. Näiteväärtus: mari.tamm@example.com"),
    ("Isik", "reg_aeg", "Isiku registreerimise kuupäev ja kellaaeg {Kohustuslik. Küsitakse süsteemi kellalt. Lubatud vahemik on 01.01.2020 00:00 kuni 31.12.2100 23:59.}. Näiteväärtus: 15.03.2026 10:30"),
    ("Isik", "viimase_muutm_aeg", "Isiku andmete viimase muutmise kuupäev ja kellaaeg {Kohustuslik. Lubatud vahemik on 01.01.2020 00:00 kuni 31.12.2100 23:59. Ei tohi olla varasem kui reg_aeg.}. Näiteväärtus: 16.03.2026 09:15"),
    ("Kasutajakonto", "parool", "Kasutajakonto parooli räsiväärtus {Kohustuslik. Ei tohi olla tühi string. Avateksti parooli ei salvestata.}. Näiteväärtus: $2b$12$examplehash"),
    ("Kasutajakonto", "on_aktiivne", "Tunnus, mis näitab, kas kontoga saab süsteemi sisse logida {Kohustuslik. Lubatud väärtused on tõene ja väär.}. Näiteväärtus: tõene"),
    ("Töötaja_roll", "kirjeldus", "Töötaja rollist tulenevate õiguste ja kohustuste kirjeldus töötajate haldurile {Võib puududa. Kui väärtus registreeritakse, siis ei tohi see olla tühi string.}. Näiteväärtus: Saab registreerida ja muuta treeninguid."),
    ("Töötaja_rolli_omamine", "alguse_aeg", "Kuupäev ja kellaaeg, millest alates töötaja roll kehtib {Kohustuslik. Lubatud vahemik on 01.01.2020 00:00 kuni 31.12.2100 23:59.}. Näiteväärtus: 01.02.2026 09:00"),
    ("Töötaja_rolli_omamine", "lopu_aeg", "Kuupäev ja kellaaeg, milleni töötaja roll kehtib {Võib puududa. Kui väärtus registreeritakse, siis peab see olema hilisem kui alguse_aeg ja lubatud vahemikus 01.01.2020 00:00 kuni 31.12.2100 23:59.}. Näiteväärtus: 31.12.2026 17:00"),
    ("Treening", "treeningu_kood", "Treeningu arvuline kood, mille järgi töötaja saab treeningut üheselt eristada {Kohustuslik. Positiivne täisarv. Unikaalne.}. Näiteväärtus: 1001"),
    ("Treening", "nimetus", "Treeningu ametlik nimetus {Kohustuslik. Ei tohi olla tühi string. Unikaalne.}. Näiteväärtus: Hommikune jooga"),
    ("Treening", "kirjeldus", "Klientidele ja töötajatele mõeldud treeningu sisu kirjeldus {Kohustuslik. Ei tohi olla tühi string.}. Näiteväärtus: Rahulik liikuvust ja hingamist arendav treening."),
    ("Treening", "kestus_minutites", "Treeningu kestus minutites {Kohustuslik. Lubatud vahemik on 15 kuni 240. Väärtus ei tohi olla 0 ega negatiivne.}. Näiteväärtus: 60"),
    ("Treening", "maksimaalne_osalejate_arv", "Treeningule lubatud maksimaalne osalejate arv inimestes {Kohustuslik. Positiivne täisarv. Väärtus ei tohi olla 0 ega negatiivne.}. Näiteväärtus: 20"),
    ("Treening", "vajalik_varustus", "Klientidele ja treenerile mõeldud kirjeldus treeningul vajaliku varustuse kohta {Kohustuslik. Ei tohi olla tühi string.}. Näiteväärtus: Joogamatt ja veepudel"),
    ("Treening", "hind", "Treeningu hind eurodes koos käibemaksuga {Kohustuslik. Null või positiivne rahasumma. Väärtus võib olla 0. Maksimaalselt kaks kohta pärast koma.}. Näiteväärtus: 12.50"),
    ("Treening", "seisund", "Treeningu elutsükli seisund {Kohustuslik. Väärtus peab vastama registreeritud Treeningu_seisundi_liik väärtusele.}. Näiteväärtus: Aktiivne"),
    ("Treening", "registreerija_e_meil", "Treeningu registreerinud kasutajakonto e-posti aadress {Kohustuslik. Peab sisaldama märki @. Peab viitama registreeritud kasutajakontole.}. Näiteväärtus: treener@example.com"),
    ("Treening", "viimase_muutja_e_meil", "Treeningut viimati muutnud kasutajakonto e-posti aadress {Kohustuslik. Peab sisaldama märki @. Peab viitama registreeritud kasutajakontole.}. Näiteväärtus: juhataja@example.com"),
    ("Treening", "reg_aeg", "Treeningu registreerimise kuupäev ja kellaaeg {Kohustuslik. Küsitakse süsteemi kellalt. Lubatud vahemik on 01.01.2020 00:00 kuni 31.12.2100 23:59.}. Näiteväärtus: 05.04.2026 12:00"),
    ("Treening", "viimase_muutm_aeg", "Treeningu viimase muutmise kuupäev ja kellaaeg {Kohustuslik. Lubatud vahemik on 01.01.2020 00:00 kuni 31.12.2100 23:59. Ei tohi olla varasem kui reg_aeg.}. Näiteväärtus: 06.04.2026 08:45"),
]


CRUD_BY_ENTITY = {
    "Klassifikaator": ["R", "R", "R", "", "R", "", "", "R", "R", "R", "R"],
    "Riik": ["R", "", "", "", "", "", "", "", "", "", ""],
    "Isiku_seisundi_liik": ["R", "", "", "", "", "", "", "", "", "", ""],
    "Töötaja_seisundi_liik": ["R", "", "", "", "", "", "", "", "", "", ""],
    "Töötaja_roll": ["R", "", "", "", "", "", "", "", "", "", ""],
    "Treeningu_seisundi_liik": ["", "R", "R", "R", "R", "R", "R", "R", "R", "R", "R"],
    "Treeningu_kategooria_tüüp": ["", "R", "R", "", "R", "", "", "R", "R", "R", "R"],
    "Treeningu_kategooria": ["", "R", "R", "", "R", "", "", "R", "R", "R", "R"],
    "Isik": ["R", "", "", "", "", "", "", "", "", "", ""],
    "Kasutajakonto": ["R", "", "", "", "", "", "", "", "", "", ""],
    "Töötaja": ["R", "", "", "", "", "", "", "", "", "", ""],
    "Töötaja_rolli_omamine": ["R", "", "", "", "", "", "", "", "", "", ""],
    "Treening": ["", "C", "RU", "RU", "RU", "RU", "RU", "R", "R", "R", "R"],
    "Treeningu_kategooria_omamine": ["", "C", "CRD", "", "R", "", "", "R", "R", "R", "R"],
}


BUSINESS_RULES = [
    "Treeningu nimetus peab olema treeningute registris unikaalne.",
    "Treeningu saab aktiivseks muuta ainult siis, kui treening kuulub vähemalt ühte treeningu kategooriasse.",
    "Aktiivne treening peab olema seotud aktiivse treeningu seisundi liigiga.",
    "Treeningu kestus peab olema 15 kuni 240 minutit.",
    "Treeningu maksimaalne osalejate arv peab olema positiivne täisarv.",
    "Treeningu hind peab olema null või positiivne rahasumma.",
    "Treeningu viimase muutmise aeg ei tohi olla varasem kui treeningu registreerimise aeg.",
]


NFR = [
    ("Kasutajaliides", "Süsteemi kasutajaliides peab olema eestikeelne ning kasutama kogu treeningute funktsionaalses allsüsteemis samu mõisteid."),
    ("Kasutajaliides", "Vormides peab kohustuslikke välju eristama enne andmete salvestamist."),
    ("Turvalisus", "Treeneri ja juhataja kasutusjuhud peavad olema kasutatavad ainult autenditud ja autoriseeritud kasutajatele."),
    ("Turvalisus", "Paroole ei tohi salvestada avatekstina."),
    ("Turvalisus", "Treeningu muutmise ja seisundimuutmise operatsioonide juures tuleb säilitada viimase muutja e-posti aadress ning muutmise aeg."),
    ("Töökindlus", "Treeningute registri varukoopia tuleb teha vähemalt üks kord ööpäevas."),
    ("Jõudlus", "Aktiivsete treeningute nimekirja päring peab tavalise koormuse korral vastama kuni kahe sekundiga."),
]


@dataclass
class CaptionCounter:
    figures: int = 0
    tables: int = 0


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "",
        "/Library/Fonts/Arial.ttf",
    ]
    for path in candidates:
        if path and Path(path).exists():
            try:
                return ImageFont.truetype(path, size=size)
            except OSError:
                pass
    return ImageFont.load_default()


def draw_centered(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, fill: str, fnt: ImageFont.ImageFont) -> None:
    lines = wrap_text(draw, text, fnt, max_width=box[2] - box[0] - 18)
    heights = [draw.textbbox((0, 0), line, font=fnt)[3] for line in lines]
    total = sum(heights) + max(0, len(lines) - 1) * 4
    y = box[1] + ((box[3] - box[1]) - total) // 2
    for line, height in zip(lines, heights):
        width = draw.textbbox((0, 0), line, font=fnt)[2]
        draw.text((box[0] + ((box[2] - box[0]) - width) // 2, y), line, fill=fill, font=fnt)
        y += height + 4


def wrap_text(draw: ImageDraw.ImageDraw, text: str, fnt: ImageFont.ImageFont, max_width: int) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        candidate = word if not current else f"{current} {word}"
        if draw.textbbox((0, 0), candidate, font=fnt)[2] <= max_width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines or [text]


def draw_box(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, fill: str = "#eef4ff", outline: str = "#1f4e79") -> None:
    draw.rounded_rectangle(box, radius=10, fill=fill, outline=outline, width=2)
    draw_centered(draw, box, text, "#1b1b1b", font(22))


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], fill: str = "#555555") -> None:
    draw.line([start, end], fill=fill, width=3)
    ex, ey = end
    sx, sy = start
    if abs(ex - sx) >= abs(ey - sy):
        direction = 1 if ex > sx else -1
        points = [(ex, ey), (ex - 14 * direction, ey - 8), (ex - 14 * direction, ey + 8)]
    else:
        direction = 1 if ey > sy else -1
        points = [(ex, ey), (ex - 8, ey - 14 * direction), (ex + 8, ey - 14 * direction)]
    draw.polygon(points, fill=fill)


def make_diagrams() -> list[tuple[Path, str]]:
    DIAGRAM_DIR.mkdir(parents=True, exist_ok=True)
    diagrams: list[tuple[Path, str]] = []

    specs = [
        ("01_architecture.png", "Jõusaali infosüsteemi põhiobjektid ja registrid", ["Isik", "Töötaja", "Treening", "Klassifikaator"], "ring"),
        ("02_use_cases.png", "Treeningute funktsionaalse allsüsteemi kasutusjuhud", ["Treener", "Registreeri", "Muuda", "Aktiveeri", "Juhataja", "Lõpeta", "Aruanne", "Klient", "Vaata aktiivseid"], "usecase"),
        ("03_activity_register.png", "Treeningu registreerimise tegevusvoog", ["Algus", "Sisesta andmed", "Kontrolli", "Salvesta", "Kuva detailandmed", "Lõpp"], "flow"),
        ("04_activity_activate.png", "Treeningu aktiveerimise tegevusvoog", ["Algus", "Vali treening", "Kontrolli kategooriat", "Muuda aktiivseks", "Lõpp"], "flow"),
        ("05_conceptual.png", "Treeningute registri kontseptuaalne eskiismudel", ["Treening", "Treeningu_seisundi_liik", "Treeningu_kategooria", "Treeningu_kategooria_tüüp", "Töötaja"], "model"),
        ("06_state.png", "Treeningu seisundidiagramm", ["Ootel", "Aktiivne", "Mitteaktiivne", "Lõppenud", "Unustatud"], "state"),
        ("07_physical.png", "Treeningute registri füüsiline andmemudel", ["treening", "treeningu_seisundi_liik", "kasutajakonto", "treeningu_kategooria_omamine", "treeningu_kategooria", "treeningu_kategooria_tyyp"], "model"),
    ]

    for filename, title, items, kind in specs:
        path = DIAGRAM_DIR / filename
        img = Image.new("RGB", (1400, 850), "#ffffff")
        draw = ImageDraw.Draw(img)
        draw.text((50, 38), title, fill="#203040", font=font(34, bold=True))
        if kind == "ring":
            coords = [(120, 230), (500, 230), (880, 230), (500, 500)] if len(items) == 4 else [(120, 170), (500, 170), (880, 170), (120, 470), (500, 470), (880, 470)]
            for (x, y), item in zip(coords, items):
                draw_box(draw, (x, y, x + 300, y + 120), item)
            if len(items) == 4:
                draw_arrow(draw, (420, 290), (500, 290))
                draw_arrow(draw, (800, 290), (880, 290))
                draw_arrow(draw, (650, 350), (650, 500))
            else:
                draw_arrow(draw, (420, 230), (500, 230))
                draw_arrow(draw, (800, 230), (880, 230))
                draw_arrow(draw, (270, 290), (270, 470))
                draw_arrow(draw, (650, 290), (650, 470))
                draw_arrow(draw, (1030, 290), (1030, 470))
        elif kind == "usecase":
            actor_boxes = [(80, 170, 290, 260), (80, 430, 290, 520), (80, 610, 290, 700)]
            usecase_boxes = [(440, 145, 710, 235), (760, 145, 1030, 235), (1080, 145, 1320, 235), (440, 410, 710, 500), (760, 410, 1030, 500), (440, 605, 710, 695)]
            for box, item in zip(actor_boxes, ["Treener", "Juhataja", "Klient ja Uudistaja"]):
                draw_box(draw, box, item, "#fff3df", "#9a5b00")
            for box, item in zip(usecase_boxes, ["Registreeri treening", "Muuda treeningut", "Aktiveeri treening", "Lõpeta treening", "Vaata aruannet", "Vaata aktiivseid"]):
                draw.ellipse(box, fill="#eef4ff", outline="#1f4e79", width=3)
                draw_centered(draw, box, item, "#1b1b1b", font(20))
            for end in [(440, 190), (760, 190), (1080, 190)]:
                draw_arrow(draw, (290, 215), end)
            for end in [(440, 455), (760, 455)]:
                draw_arrow(draw, (290, 475), end)
            draw_arrow(draw, (290, 655), (440, 650))
        elif kind == "state":
            coords = [(105, 340), (430, 180), (430, 500), (830, 180), (830, 500)]
            for (x, y), item in zip(coords, items):
                draw_box(draw, (x, y, x + 230, y + 105), item, "#f2f7ed", "#3d6b2f")
            draw_arrow(draw, (335, 392), (430, 232))
            draw_arrow(draw, (335, 392), (830, 552))
            draw_arrow(draw, (545, 285), (545, 500))
            draw_arrow(draw, (615, 500), (615, 285))
            draw_arrow(draw, (660, 232), (830, 232))
            draw_arrow(draw, (660, 552), (830, 232))
        elif kind == "flow":
            x = 70
            box_width = 170
            step_width = 220
            boxes = []
            for item in items:
                boxes.append((x, 320, x + box_width, 430, item))
                x += step_width
            for left, top, right, bottom, item in boxes:
                draw_box(draw, (left, top, right, bottom), item, "#f8fbff", "#1f4e79")
            for first, second in zip(boxes, boxes[1:]):
                draw_arrow(draw, (first[2], 375), (second[0], 375))
        else:
            if "treeningu_kategooria_omamine" in items:
                boxes = {
                    "treening": (70, 335, 330, 455),
                    "treeningu_seisundi_liik": (430, 160, 740, 280),
                    "kasutajakonto": (430, 520, 740, 640),
                    "treeningu_kategooria_omamine": (790, 335, 1100, 455),
                    "treeningu_kategooria": (1050, 160, 1360, 280),
                    "treeningu_kategooria_tyyp": (1050, 520, 1360, 640),
                }
                for item in items:
                    draw_box(draw, boxes[item], item)
                draw_arrow(draw, (330, 365), (430, 220))
                draw_arrow(draw, (330, 425), (430, 580))
                draw_arrow(draw, (330, 395), (790, 395))
                draw_arrow(draw, (1100, 395), (1205, 280))
                draw_arrow(draw, (1205, 280), (1205, 520))
            else:
                boxes = {
                    "Treening": (90, 335, 380, 455),
                    "Treeningu_seisundi_liik": (500, 160, 850, 280),
                    "Töötaja": (930, 335, 1220, 455),
                    "Treeningu_kategooria": (500, 520, 850, 640),
                    "Treeningu_kategooria_tüüp": (930, 520, 1220, 640),
                }
                for item in items:
                    draw_box(draw, boxes[item], item)
                draw_arrow(draw, (380, 365), (500, 220))
                draw_arrow(draw, (380, 395), (930, 395))
                draw_arrow(draw, (380, 425), (500, 580))
                draw_arrow(draw, (850, 580), (930, 580))
        img.save(path)
        diagrams.append((path, title))
    return diagrams


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def style_table(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header = table.rows[0]
    set_repeat_table_header(header)
    for cell in header.cells:
        set_cell_shading(cell, "D9EAF7")
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def shrink_table_text(table, size: int) -> None:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(size)


def add_table(doc: Document, caption: str, headers: list[str], rows: list[Iterable[str]], counter: CaptionCounter):
    counter.tables += 1
    p = doc.add_paragraph(f"Tabel {counter.tables}. {caption}", style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table = doc.add_table(rows=1, cols=len(headers))
    style_table(table)
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for row_data in rows:
        row = table.add_row()
        for i, value in enumerate(row_data):
            row.cells[i].text = str(value)
    doc.add_paragraph()
    return table


def crud_summary(values: list[str]) -> str:
    ordered = "CRUD"
    present = {letter for value in values for letter in value if letter in ordered}
    return "".join(letter for letter in ordered if letter in present)


def add_figure(doc: Document, image_path: Path, caption: str, counter: CaptionCounter) -> None:
    counter.figures += 1
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(6.6))
    cp = doc.add_paragraph(f"Joonis {counter.figures}. {caption}", style="Caption")
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER


def bullet_list(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def number_list(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_heading(doc: Document, text: str, level: int) -> None:
    doc.add_heading(text, level=level)


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    styles["Normal"].font.size = Pt(10.5)
    for style_name, size in [("Title", 20), ("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)]:
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(31, 78, 121)
    styles["Caption"].font.name = "Arial"
    styles["Caption"].font.size = Pt(9)
    styles["Caption"].font.italic = True
    return doc


def add_title_page(doc: Document) -> None:
    for value in [UNIVERSITY, FACULTY, INSTITUTE]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(value)
    for _ in range(5):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"{SYSTEM_NAME}\n{SUBSYSTEM_NAME.lower()}")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(31, 78, 121)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(COURSE).bold = True
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(
        f"Üliõpilased: {AUTHORS}\n"
        f"Õpperühm: {STUDY_GROUP}\n"
        f"Matrikli nr: {MATRICULATION_NUMBERS}\n"
        f"e-posti aadressid: {AUTHOR_EMAILS}\n"
        f"Juhendaja: {SUPERVISOR}\n"
        "Tallinn 2026"
    )
    doc.add_page_break()


def add_contents(doc: Document, counter: CaptionCounter) -> None:
    add_heading(doc, "Sisukord", 1)
    doc.add_paragraph("Sisukord on esitatud dokumendi struktuurse ülevaatena. Leheküljenumbreid ei hoita käsitsi, et vältida aegunud viiteid pärast dokumendi taastootmist.")
    rows = [
        ("1", "Sissejuhatus"),
        ("2", "Süsteemi üldvaade"),
        ("3", "Mittefunktsionaalsed nõuded"),
        ("4", "Kasutusjuhud"),
        ("5", "Treeningute registri eskiismudel"),
        ("6", "Operatsioonilepingud"),
        ("7", "Andmemudel ja füüsiline disain"),
        ("8", "Reprodutseerimine ja kontroll"),
    ]
    add_table(doc, "Dokumendi peatükkide ülevaade", ["Peatükk", "Pealkiri"], rows, counter)
    doc.add_page_break()


def add_overview(doc: Document, counter: CaptionCounter, diagrams: list[tuple[Path, str]]) -> None:
    add_heading(doc, "1 Sissejuhatus", 1)
    doc.add_paragraph("Dokument kirjeldab jõusaali infosüsteemi treeningute funktsionaalset allsüsteemi ning selle allsüsteemi toetavat treeningute registrit.")
    add_heading(doc, "1.1 Organisatsiooni kirjeldus", 2)
    doc.add_paragraph("Jõusaal pakub klientidele rühmatreeninguid, personaaltreeninguid ja muid treeninguteenuseid. Organisatsiooni igapäevane töö sõltub sellest, et treeningute info oleks täpne, kontrollitud ja klientidele õigel ajal kättesaadav.")
    add_heading(doc, "1.2 Organisatsiooni eesmärgid", 2)
    bullet_list(doc, [
        "Pakkuda klientidele kvaliteetseid ja ajakohase kavaga treeninguteenuseid.",
        "Suurendada klientide järjepidevat treeningutes osalemist.",
        "Toetada treenerite töö planeerimist ja treeningute portfelli juhtimist.",
        "Võimaldada juhtkonnal hinnata treeningute valikut ja selle muutumist.",
    ])

    add_heading(doc, "2 Süsteemi üldvaade", 1)
    doc.add_paragraph("Süsteemi üldvaade seob organisatsiooni eesmärgid põhiobjektide, pädevusalade, registrite ja põhiprotsessidega.")
    add_figure(doc, diagrams[0][0], diagrams[0][1], counter)

    add_heading(doc, "2.1 Infosüsteemi eesmärgid", 2)
    bullet_list(doc, [
        "Võimaldada treeningut elektrooniliselt registreerida, muuta ja seisundite kaudu juhtida.",
        "Säilitada isikute ja töötajate andmeid mahus, mis võimaldab tööalaseid õigusi ja vastutust üheselt tuvastada.",
        "Säilitada treeningute kohta andmeid mahus, mis toetab treeningute pakkumist ja aruandlust.",
        "Tagada aktiivsete treeningute ajakohane avalik kuvamine klientidele ja uudistajatele.",
        "Võimaldada klassifikaatorite abil ühtlustada treeningute seisundeid, kategooriaid ja töötajatega seotud liigitusi.",
        "Võimaldada juhtkonnal vaadata treeningute koondaruannet seisundite ja kategooriate kaupa.",
    ])

    add_heading(doc, "2.2 Põhiobjektid", 2)
    add_table(doc, "Põhiobjektid", ["Põhiobjekt", "Selgitus"], MAIN_OBJECTS, counter)

    add_heading(doc, "2.3 Tegutsejad ja pädevusalad", 2)
    add_table(doc, "Tegutsejad ja pädevusalad", ["Tegutseja", "Pädevusala liik", "Huvi"], ACTORS, counter)

    add_heading(doc, "2.4 Funktsionaalsed allsüsteemid ja registrid", 2)
    add_table(doc, "Funktsionaalsete allsüsteemide ja registrite vastavus", ["Funktsionaalne allsüsteem", "Register", "Põhiobjekt"], SUBSYSTEMS_REGISTERS, counter)

    add_heading(doc, "2.5 Põhiprotsessid ja käivitavad sündmused", 2)
    add_table(doc, "Põhiprotsessid", ["Põhiprotsess", "Käivitav sündmus", "Tulemus"], PROCESSES, counter)

    add_heading(doc, "2.6 Lausendid", 2)
    number_list(doc, STATEMENTS)


def add_nfr(doc: Document, counter: CaptionCounter) -> None:
    add_heading(doc, "3 Mittefunktsionaalsed nõuded", 1)
    doc.add_paragraph("Mittefunktsionaalsed nõuded kirjeldavad süsteemi kvaliteediomadusi ja tehnilisi piiranguid, mitte eraldiseisvaid ärifunktsioone.")
    add_table(doc, "Mittefunktsionaalsed nõuded", ["Alajaotus", "Nõue"], NFR, counter)


def add_use_cases(doc: Document, counter: CaptionCounter, diagrams: list[tuple[Path, str]]) -> None:
    add_heading(doc, "4 Kasutusjuhud", 1)
    doc.add_paragraph("Kasutusjuhud kirjeldavad treeningute funktsionaalse allsüsteemi kasutajale nähtavat käitumist ja süsteemi vastuseid.")
    add_figure(doc, diagrams[1][0], diagrams[1][1], counter)

    add_heading(doc, "4.1 Kõrgtaseme kasutusjuhud", 2)
    add_table(doc, "Kõrgtaseme kasutusjuhud", ["Kasutusjuht", "Tegutsejad", "Kirjeldus"], HIGH_LEVEL_USE_CASES, counter)

    add_heading(doc, "4.2 Laiendatud kasutusjuhud", 2)
    doc.add_paragraph("Jaotis sisaldab tabelis 7 esitatud kõrgtaseme kasutusjuhtude laiendatud kirjeldusi. Iga kirjeldus esitab primaarse tegutseja, osapooled ja huvid, eeltingimused, tüüpilise sündmuste järjestuse, alternatiivsed käigud ning järeltingimused.")
    for index, uc in enumerate(EXTENDED_USE_CASES, start=1):
        add_heading(doc, f"4.2.{index} {uc['name']}", 3)
        add_table(doc, f"Laiendatud kasutusjuht: {uc['name']}", ["Väli", "Sisu"], [
            ("Kasutusjuht", uc["name"]),
            ("Primaarne tegutseja", uc["actor"]),
            ("Käivitav sündmus", uc["trigger"]),
            ("Eeltingimused", "\n".join(uc["pre"])),
            ("Järeltingimused", "\n".join(uc["post"])),
        ], counter)
        add_table(doc, f"Osapooled ja nende huvid: {uc['name']}", ["Osapool", "Huvi"], uc["stakeholders"], counter)
        add_table(doc, f"Stsenaarium: {uc['name']}", ["Samm", "Tegevus", "Operatsioon"], [(i + 1, step, op) for i, (step, op) in enumerate(uc["steps"])], counter)
        add_table(doc, f"Laiendused: {uc['name']}", ["Laiendus"], [(item,) for item in uc["extensions"]], counter)

    add_figure(doc, diagrams[2][0], diagrams[2][1], counter)
    add_figure(doc, diagrams[3][0], diagrams[3][1], counter)


def add_register_model(doc: Document, counter: CaptionCounter, diagrams: list[tuple[Path, str]]) -> None:
    add_heading(doc, "5 Treeningute registri eskiismudel", 1)
    doc.add_paragraph("Treeningute register on treeningute funktsionaalse allsüsteemi põhiregister. Register kasutab töötajate, isikute ja klassifikaatorite registrites olevaid andmeid.")
    add_heading(doc, "5.1 Registri seosed", 2)
    add_table(doc, "Treeningute registri eskiismudeli seosed", ["Nimekiri", "Väärtused"], [
        ("Registrit kasutavad pädevusalad", "Treener; Juhataja; Klient; Uudistaja"),
        ("Registrit teenindavad funktsionaalsed allsüsteemid", "Treeningute funktsionaalne allsüsteem"),
        ("Selle registriga seotud registrid", "Isikute register; Töötajate register; Klassifikaatorite register"),
    ], counter)
    add_heading(doc, "5.2 Ärireeglid", 2)
    number_list(doc, BUSINESS_RULES)
    add_heading(doc, "5.3 Kontseptuaalne eskiismudel", 2)
    add_figure(doc, diagrams[4][0], diagrams[4][1], counter)


def add_operations(doc: Document, counter: CaptionCounter) -> None:
    add_heading(doc, "6 Operatsioonilepingud", 1)
    doc.add_paragraph("Operatsioonilepingud on koostatud lepingprojekteerimise põhimõttel ning nende tähised vastavad kasutusjuhtudes kasutatud viidetele.")
    for op_id, name, pre, post in OPERATIONS:
        add_heading(doc, f"{op_id} {name}", 2)
        add_table(doc, f"Andmebaasioperatsiooni leping: {op_id}", ["Osa", "Sisu"], [
            ("Operatsioon", f"{op_id} {name}"),
            ("Eeltingimused", "\n".join(pre)),
            ("Järeltingimused", "\n".join(post)),
            ("Kasutus kasutusjuhtude poolt", ", ".join(sorted({uc["name"] for uc in EXTENDED_USE_CASES if any(op_id == op for _, op in uc["steps"])})) or "Tööalane abipäring"),
        ], counter)


def add_data_design(doc: Document, counter: CaptionCounter, diagrams: list[tuple[Path, str]]) -> None:
    add_heading(doc, "7 Andmemudel ja füüsiline disain", 1)
    doc.add_paragraph("Andmemudel kirjeldab treeningute registri loogilisi olemitüüpe ning PostgreSQL füüsilist realisatsiooni.")
    add_heading(doc, "7.1 Olemid ja atribuudid", 2)
    add_table(doc, "Olemitüübid", ["Olemitüüp", "Definitsioon", "Identifikaator"], ENTITIES, counter)
    add_table(doc, "Atribuutide definitsioonid", ["Olemitüüp", "Atribuut", "Definitsioon"], ATTRIBUTES, counter)
    use_case_names = [name for name, _, _ in HIGH_LEVEL_USE_CASES]
    crud_headers = ["Olemitüüp", *use_case_names, "Kokku"]
    crud_rows = []
    for entity, *_ in ENTITIES:
        values = CRUD_BY_ENTITY[entity]
        crud_rows.append([entity, *values, crud_summary(values)])
    crud_table = add_table(doc, "CRUD maatriks", crud_headers, crud_rows, counter)
    shrink_table_text(crud_table, 6)
    add_heading(doc, "7.2 Seisundimudel", 2)
    add_figure(doc, diagrams[5][0], diagrams[5][1], counter)
    add_heading(doc, "7.3 Füüsiline mudel", 2)
    add_figure(doc, diagrams[6][0], diagrams[6][1], counter)
    add_heading(doc, "7.4 SQL DDL", 2)
    doc.add_paragraph("Järgmine DDL on PostgreSQL jaoks koostatud ning seda kasutatakse projekti kontrollis.")
    for line in SQL_DDL.strip().splitlines():
        display_line = line.replace("REFERENCES", "references")
        p = doc.add_paragraph()
        run = p.add_run(display_line)
        run.font.name = "Courier New"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
        run.font.size = Pt(8)


def add_reproducibility(doc: Document, counter: CaptionCounter) -> None:
    add_heading(doc, "8 Reprodutseerimine ja kontroll", 1)
    doc.add_paragraph("Projekt on taastoodetav jälgitavatest lähtefailidest. Lõppartefaktid genereeritakse skriptiga, mis loob DOCX dokumendi, teisendab EAP malli, rakendab EAP parandused ning väljastab PostgreSQL DDL skripti.")
    add_table(doc, "Taastootmise sisendid ja väljundid", ["Fail või kataloog", "Roll"], [
        ("instruction_guides/", "Kursuse juhendid, näidisprojekt ja tüüpvigade dokument."),
        ("preset_files/EA_mall_AB_projekt_Eeltaidetud_2026.eap", "EAP lähtebaas, millest lõplik mudel teisendatakse."),
        ("tools/fill_report_docx.py", "Struktureeritud DOCX dokumendi generaator."),
        ("tools/EapConvert.java", "EAP malli teisendaja kirjutatavasse Access 2000 vormingusse."),
        ("tools/EapRename.java ja tools/EapFixes.java", "EAP mudeli nime- ja sisuparandused."),
        ("tools/sql_ddl.py", "PostgreSQL DDL lähtefail."),
        ("jousaali_skript.sql", "Esitamiseks mõeldud PostgreSQL tabelite loomise skript."),
        ("rakendus/", "Treeneri ja juhataja töökoha prototüübi lähtekood ja käivitamisjuhend."),
        ("tools/validate_project.py", "Automaatne kontrollskript."),
        ("build_all.sh või build_all.bat", "Lõppartefaktide taastootmine puhtast kloonist."),
    ], counter)


def validate_text_sanity(doc: Document) -> None:
    text = "\n".join(p.text for p in doc.paragraphs)
    forbidden = ["HYPERLINK", "PAGEREF", "MERGEFORMAT", "<täienda", "<Siia", "<abc>"]
    found = [item for item in forbidden if item in text]
    if found:
        raise ValueError(f"Generated DOCX contains forbidden visible text: {found}")
    if re.search(r"\bOP\d+(?:\.\d+)?\b", text):
        refs = set(re.findall(r"\bOP\d+(?:\.\d+)?\b", text))
        defs = {op_id for op_id, *_ in OPERATIONS}
        dangling = sorted(refs - defs)
        if dangling:
            raise ValueError(f"Dangling operation references: {dangling}")


def main() -> None:
    diagrams = make_diagrams()
    counter = CaptionCounter()
    doc = setup_document()
    add_title_page(doc)
    add_contents(doc, counter)
    add_overview(doc, counter, diagrams)
    add_nfr(doc, counter)
    add_use_cases(doc, counter, diagrams)
    add_register_model(doc, counter, diagrams)
    add_operations(doc, counter)
    add_data_design(doc, counter, diagrams)
    add_reproducibility(doc, counter)
    validate_text_sanity(doc)
    doc.save(DST)
    print(f"Wrote {DST}")
    print(f"Embedded figures: {counter.figures}")
    print(f"Word tables: {counter.tables}")


if __name__ == "__main__":
    main()
