#!/usr/bin/env python3
from __future__ import annotations

import csv
import io
import py_compile
import re
import shutil
import subprocess
import sys
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from PIL import Image

ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "Jousaali_infosusteemi_treeningute_funktsionaalne_allsusteem.docx"
EAP = ROOT / "Jousaali_infosusteemi_treeningute_funktsionaalne_allsusteem.eap"
GUIDES = ROOT / "instruction_guides"
SQL_OUTPUT = ROOT / "jousaali_skript.sql"
APP_DIR = ROOT / "rakendus"
sys.path.insert(0, str(ROOT / "tools"))
from sql_ddl import SQL_DDL  # noqa: E402


EXPECTED_ACTORS = {
    "Treener",
    "Juhataja",
    "Klient",
    "Uudistaja",
    "Klassifikaatorite haldur",
    "Töötajate haldur",
    "Maksekeskus",
}


def fail(message: str, failures: list[str]) -> None:
    failures.append(message)
    print(f"FAIL: {message}")


def ok(message: str) -> None:
    print(f"PASS: {message}")


def warn(message: str) -> None:
    print(f"WARN: {message}")


def docx_text(doc: Document) -> str:
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(p.text for p in cell.paragraphs)
    return "\n".join(parts)


def validate_docx(failures: list[str]) -> None:
    doc = Document(DOCX)
    text = docx_text(doc)
    section = doc.sections[0]
    width_cm = section.page_width.cm
    height_cm = section.page_height.cm
    heading_count = sum(1 for p in doc.paragraphs if p.style.name.startswith("Heading"))
    table_count = len(doc.tables)
    image_count = len(doc.inline_shapes)
    figure_captions = re.findall(r"^Joonis\s+\d+\.", text, re.M)
    table_captions = re.findall(r"^Tabel\s+\d+\.", text, re.M)

    if abs(width_cm - 21.0) <= 0.1 and abs(height_cm - 29.7) <= 0.1:
        ok(f"DOCX page size is A4 ({width_cm:.1f} x {height_cm:.1f} cm)")
    else:
        fail(f"DOCX page size is not A4 ({width_cm:.1f} x {height_cm:.1f} cm)", failures)
    if heading_count >= 20:
        ok(f"DOCX has real heading styles ({heading_count})")
    else:
        fail(f"DOCX heading style count too low ({heading_count})", failures)
    if table_count >= 10:
        ok(f"DOCX has real Word tables ({table_count})")
    else:
        fail(f"DOCX table count too low ({table_count})", failures)
    if image_count >= 7:
        ok(f"DOCX has embedded images/diagrams ({image_count})")
    else:
        fail(f"DOCX embedded image count too low ({image_count})", failures)
    if len(figure_captions) >= image_count and len(table_captions) >= table_count:
        ok(f"DOCX captions exist for figures ({len(figure_captions)}) and tables ({len(table_captions)})")
    else:
        fail(f"DOCX caption counts are insufficient: figures={len(figure_captions)}, tables={len(table_captions)}", failures)

    forbidden = ["HYPERLINK", "PAGEREF", "REF ", "SEQ", "MERGEFORMAT", "<täienda>", "<täienda või kustuta>", "<Siia"]
    found = [item for item in forbidden if item in text]
    if found:
        fail(f"DOCX contains forbidden visible field/placeholder text: {found}", failures)
    else:
        ok("DOCX has no visible forbidden field-code or placeholder text")

    known_bad_phrases = [
        "treeningupäevikute register",
        "treening andmeid",
        "treening kategooriasse",
        "treening koondaruanne",
        "treening arv",
        "\ntreener ",
    ]
    found_bad = [phrase for phrase in known_bad_phrases if phrase in text]
    if found_bad:
        fail(f"DOCX contains known bad wording: {found_bad}", failures)
    else:
        ok("DOCX known bad wording is absent")

    required_title_info = [
        "Tristan Aik Sild",
        "Gustav Tamkivi",
        "Õpperühm: IAIB23",
        "253782IAIB",
        "253787IAIB",
        "gustav@taltech.ee",
        "trists@taltech.ee",
    ]
    missing_title_info = [item for item in required_title_info if item not in text]
    if missing_title_info:
        fail(f"DOCX title page is missing author/student information: {missing_title_info}", failures)
    else:
        ok("DOCX title page contains author names, study group, matriculation numbers, and emails")

    op_refs = set(re.findall(r"\bOP\d+(?:\.\d+)?\b", text))
    op_defs = set(re.findall(r"^OP\d+\s+", text, re.M))
    op_defs = {value.strip() for value in op_defs}
    dangling = sorted(op_refs - op_defs)
    if dangling:
        fail(f"DOCX has dangling operation references: {dangling}", failures)
    else:
        ok(f"DOCX operation references all have definitions ({len(op_refs)} references)")

    known_bad_ops = {"OP1.1", "OP2.1", "OP3.1", "OP8.2", "OP9.1", "OP10.1", "OP11.2", "OP11.3"}
    bad_ops_present = sorted(op_refs & known_bad_ops)
    if bad_ops_present:
        fail(f"DOCX contains obsolete operation IDs: {bad_ops_present}", failures)

    scenario_refs: set[str] = set()
    for table in doc.tables:
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        if headers == ["Samm", "Tegevus", "Operatsioon"]:
            for row in table.rows[1:]:
                scenario_refs.update(re.findall(r"\bOP\d+(?:\.\d+)?\b", row.cells[2].text))
    unused_defs = sorted(op_defs - scenario_refs, key=lambda value: int(value[2:]) if value[2:].isdigit() else 9999)
    if unused_defs:
        warn(f"DOCX defines operations not referenced from scenario operation cells: {unused_defs}")

    clipped_images = []
    with ZipFile(DOCX) as archive:
        media_names = sorted(name for name in archive.namelist() if name.startswith("word/media/") and name.lower().endswith((".png", ".jpg", ".jpeg")))
        for name in media_names:
            image = Image.open(io.BytesIO(archive.read(name))).convert("RGB")
            width, height = image.size
            edge = 12
            edge_crops = {
                "left": image.crop((0, 0, edge, height)),
                "right": image.crop((width - edge, 0, width, height)),
                "top": image.crop((0, 0, width, edge)),
                "bottom": image.crop((0, height - edge, width, height)),
            }
            for side, crop in edge_crops.items():
                nonwhite = sum(1 for pixel in crop.getdata() if pixel[0] < 245 or pixel[1] < 245 or pixel[2] < 245)
                if nonwhite > 0:
                    clipped_images.append(f"{name}:{side}")
    if clipped_images:
        fail(f"DOCX embedded diagram content touches image edge, possible clipping: {clipped_images}", failures)
    else:
        ok("DOCX embedded diagrams have clear image-edge margins")


def table_block(sql: str, table_name: str) -> str:
    match = re.search(rf"CREATE TABLE {table_name}\s*\((.*?)\n\);", sql, re.S)
    return match.group(1) if match else ""


def column_line(block: str, column: str) -> str:
    for line in block.splitlines():
        stripped = line.strip().rstrip(",")
        if stripped.startswith(column + " "):
            return stripped
    return ""


def validate_sql(failures: list[str]) -> None:
    sql = SQL_DDL
    isik = table_block(sql, "isik")
    treening = table_block(sql, "treening")
    required = [
        ("isik", isik, "synni_kp"),
        ("isik", isik, "e_meil"),
        ("treening", treening, "kirjeldus"),
        ("treening", treening, "kestus_minutites"),
        ("treening", treening, "maksimaalne_osalejate_arv"),
        ("treening", treening, "vajalik_varustus"),
        ("treening", treening, "hind"),
    ]
    missing_not_null = [f"{table}.{col}" for table, block, col in required if "NOT NULL" not in column_line(block, col)]
    if missing_not_null:
        fail(f"SQL required fields are nullable: {missing_not_null}", failures)
    else:
        ok("SQL required fields are NOT NULL")

    sql_checks = {
        "isik.e_meil unique": "CONSTRAINT uq_isik_e_meil UNIQUE (e_meil)",
        "treening.nimetus unique": "CONSTRAINT uq_treening_nimetus UNIQUE (nimetus)",
        "non-empty text checks": "btrim(nimetus) <> ''",
        "duration range": "kestus_minutites BETWEEN 15 AND 240",
        "positive participant count": "maksimaalne_osalejate_arv > 0",
        "timestamp consistency": "viimase_muutm_aeg >= reg_aeg",
    }
    missing = [name for name, needle in sql_checks.items() if needle not in sql]
    if missing:
        fail(f"SQL missing declared constraints: {missing}", failures)
    else:
        ok("SQL declared uniqueness and check constraints exist")

    keys: dict[str, set[tuple[str, ...]]] = {}
    for table in re.findall(r"CREATE TABLE\s+(\w+)\s*\(", sql):
        block = table_block(sql, table)
        table_keys: set[tuple[str, ...]] = set()
        for kind in ["PRIMARY KEY", "UNIQUE"]:
            for match in re.findall(rf"{kind}\s*\(([^)]+)\)", block):
                table_keys.add(tuple(part.strip() for part in match.split(",")))
        keys[table] = table_keys
    bad_fk = []
    for table in re.findall(r"CREATE TABLE\s+(\w+)\s*\(", sql):
        block = table_block(sql, table)
        for local_cols, target_table, target_cols in re.findall(r"FOREIGN KEY\s*\(([^)]+)\)\s*REFERENCES\s+(\w+)\s*\(([^)]+)\)", block, re.S):
            target_tuple = tuple(part.strip() for part in target_cols.split(","))
            if target_tuple not in keys.get(target_table, set()):
                bad_fk.append(f"{table}({local_cols}) -> {target_table}({target_cols})")
    if bad_fk:
        fail(f"SQL foreign keys reference non-key columns: {bad_fk}", failures)
    else:
        ok("SQL foreign keys reference primary or unique keys")


def mdb_export(table: str) -> list[dict[str, str]] | None:
    if not shutil.which("mdb-export"):
        return None
    result = subprocess.run(["mdb-export", str(EAP), table], check=True, text=True, capture_output=True)
    return list(csv.DictReader(io.StringIO(result.stdout)))


def validate_eap(failures: list[str]) -> None:
    objects = mdb_export("t_object")
    connectors = mdb_export("t_connector")
    attributes = mdb_export("t_attribute")
    if objects is None or connectors is None or attributes is None:
        warn("mdb-export is unavailable; EAP table checks were skipped")
        return

    duplicate_connectors = sorted({row["Connector_ID"] for row in connectors if sum(1 for other in connectors if other["Connector_ID"] == row["Connector_ID"]) > 1})
    duplicate_attributes = sorted({row["ID"] for row in attributes if sum(1 for other in attributes if other["ID"] == row["ID"]) > 1})
    if duplicate_connectors:
        fail(f"EAP duplicate connector IDs remain: {duplicate_connectors}", failures)
    else:
        ok("EAP has no duplicate connector IDs")
    if duplicate_attributes:
        fail(f"EAP duplicate attribute IDs remain: {duplicate_attributes}", failures)
    else:
        ok("EAP has no duplicate attribute IDs")

    stale_values = ["<Siia", "OP..", "OP ...", "<täienda", "<abc>"]
    stale = []
    for row in connectors + objects + attributes:
        combined = " ".join(row.values())
        for value in stale_values:
            if value in combined:
                stale.append(value)
    if stale:
        fail(f"EAP stale placeholders remain: {sorted(set(stale))}", failures)
    else:
        ok("EAP stale placeholders are absent")

    actors = {row["Name"] for row in objects if row["Object_Type"] == "Actor"}
    if actors == EXPECTED_ACTORS:
        ok("EAP actor set matches DOCX")
    else:
        fail(f"EAP actor mismatch. expected={sorted(EXPECTED_ACTORS)}, actual={sorted(actors)}", failures)

    unsupported = []
    for row in connectors:
        if row["Connector_Type"] == "UseCase" and row["Stereotype"].strip() == "extend":
            pair = (row["Start_Object_ID"], row["End_Object_ID"])
            if pair in {("63", "23"), ("22", "28")}:
                unsupported.append(pair)
    if unsupported:
        fail(f"EAP unsupported extend relationships remain: {unsupported}", failures)
    else:
        ok("EAP unsupported use-case extend relationships are absent")

    blank_use_cases = [row["Name"] for row in objects if row["Object_Type"] == "UseCase" and not row["Note"].strip()]
    if blank_use_cases:
        fail(f"EAP use cases with empty notes remain: {blank_use_cases}", failures)
    else:
        ok("EAP use cases have meaningful notes")

    defined_ops = set(re.findall(r"\bOP\d+\b", SQL_DDL))
    doc = Document(DOCX)
    doc_text = docx_text(doc)
    defined_ops.update(re.findall(r"^OP\d+\s+", doc_text, re.M))
    defined_ops = {value.strip() for value in defined_ops}
    connector_ops = set()
    for row in connectors:
        connector_ops.update(re.findall(r"\bOP\d+(?:\.\d+)?\b", row.get("PDATA3", "")))
    undefined_connector_ops = sorted(connector_ops - defined_ops)
    if undefined_connector_ops:
        fail(f"EAP connector operation references are not defined in DOCX: {undefined_connector_ops}", failures)
    else:
        ok("EAP connector operation references match DOCX operation definitions")

    sql_tables = set(re.findall(r"CREATE TABLE\s+(\w+)", SQL_DDL))
    eap_table_classes = {row["Name"] for row in objects if row.get("Stereotype", "").strip() == "table"}
    if sql_tables == eap_table_classes:
        ok("EAP physical table classes match SQL DDL tables")
    else:
        fail(f"EAP/SQL table mismatch. missing_in_eap={sorted(sql_tables - eap_table_classes)}, extra_in_eap={sorted(eap_table_classes - sql_tables)}", failures)


def validate_repo(failures: list[str]) -> None:
    expected_guides = {
        "Iseseisva_too_ylesande_pystitus_ITI0206_2026.pdf",
        "Projekti_mustripohine_juhend_1_52.pdf",
        "Projekti_tyypvead_ITI0206_2026.pdf",
        "Naidisprojekt_ITI0206_vastuvotuajad_ver6_44.pdf",
    }
    actual = {path.name for path in GUIDES.glob("*.pdf")}
    missing = sorted(expected_guides - actual)
    if missing:
        fail(f"instruction_guides is missing PDFs: {missing}", failures)
    else:
        ok("instruction_guides contains required PDFs")
    for path in [
        "build_all.sh",
        "build_all.bat",
        "requirements.txt",
        "preset_files/EA_converted_source.eap",
        "tools/fill_report_docx.py",
        "tools/EapFixes.java",
        "jousaali_skript.sql",
        "rakendus/app.py",
        "rakendus/requirements.txt",
        "rakendus/.env.example",
        "rakendus/README.md",
        "rakendus/SETUP.sh",
        "rakendus/test_data.sql",
        "rakendus/templates/base.html",
        "rakendus/templates/dashboard.html",
        "rakendus/templates/error.html",
        "rakendus/templates/login.html",
        "rakendus/templates/register_training.html",
        "rakendus/templates/trainings.html",
    ]:
        if not (ROOT / path).exists():
            fail(f"required reproducibility file is missing: {path}", failures)
    ok("required reproducibility files are present")

    if SQL_OUTPUT.exists():
        generated_sql = SQL_DDL.strip()
        file_sql = SQL_OUTPUT.read_text(encoding="utf-8").strip()
        if file_sql == generated_sql:
            ok("standalone SQL script matches tools/sql_ddl.py")
        else:
            fail("standalone SQL script does not match tools/sql_ddl.py output", failures)

    if APP_DIR.exists():
        try:
            py_compile.compile(str(APP_DIR / "app.py"), doraise=True)
            ok("application prototype Python source compiles")
        except py_compile.PyCompileError as exc:
            fail(f"application prototype Python source does not compile: {exc}", failures)

        test_data_path = APP_DIR / "test_data.sql"
        test_data = test_data_path.read_text(encoding="utf-8") if test_data_path.exists() else ""
        if re.search(r"[\u0400-\u04FF]", test_data):
            fail("application test_data.sql contains Cyrillic-looking corrupted characters", failures)
        else:
            ok("application test_data.sql has no Cyrillic-looking corrupted characters")

        required_seed_terms = ["treeningu_seisundi_liik", "treeningu_kategooria_tyyp", "treeningu_kategooria"]
        missing_seed_terms = [term for term in required_seed_terms if term not in test_data]
        if missing_seed_terms:
            fail(f"application test_data.sql is missing required seed data terms: {missing_seed_terms}", failures)
        else:
            ok("application test_data.sql contains required classifier seed data")

    local_only = [path for path in [APP_DIR / ".env", APP_DIR / "venv"] if path.exists()]
    if local_only:
        fail(f"application contains local-only files that should not be committed: {[str(path.relative_to(ROOT)) for path in local_only]}", failures)


def main() -> int:
    failures: list[str] = []
    validate_docx(failures)
    validate_sql(failures)
    validate_eap(failures)
    validate_repo(failures)
    if failures:
        print(f"\nValidation failed with {len(failures)} issue(s).")
        return 1
    print("\nValidation passed.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
